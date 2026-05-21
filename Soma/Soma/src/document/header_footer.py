"""
Header & footer setup.

Adds a placeholder header (room for a future logo/title) and a footer
with an auto-updating page-number field. The section margins are sized
to reserve 1.25" for both header and footer per the BPD spec.
"""

from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HEADER_PLACEHOLDER = "[Company Logo / Document Title]"
FOOTER_PLACEHOLDER = "Business Process Document"


def _add_page_number_field(paragraph) -> None:
    """Append a {PAGE} field to the given paragraph."""
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(end)


def configure_page_layout(doc: DocxDocument) -> None:
    """Set section margins so header/footer have 1.25" of room."""
    for section in doc.sections:
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(1.25)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)


def add_header_and_footer(doc: DocxDocument) -> None:
    """Populate every section's header and footer."""
    for section in doc.sections:
        # Header — placeholder paragraph users can replace with a logo.
        header = section.header
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_run = h_para.add_run(HEADER_PLACEHOLDER)
        h_run.font.size = Pt(9)
        h_run.font.italic = True

        # Footer — left-aligned placeholder + right-aligned page number.
        footer = section.footer
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        f_run = f_para.add_run(f"{FOOTER_PLACEHOLDER}\t\tPage ")
        f_run.font.size = Pt(9)
        _add_page_number_field(f_para)
